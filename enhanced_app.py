import streamlit as st
import plotly.express as px
import plotly.graph_objects as go
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import RetrievalQA, ConversationalRetrievalChain
from langchain.chat_models import ChatOpenAI
from langchain.memory import ConversationBufferWindowMemory
from langchain.prompts import PromptTemplate
import os
import json
import pandas as pd
from datetime import datetime, timedelta
import hashlib
import pickle
from pathlib import Path
import re
import time
from io import BytesIO
try:
    import docx
except ImportError:
    docx = None
from PIL import Image
try:
    from wordcloud import WordCloud
    import matplotlib.pyplot as plt
except ImportError:
    WordCloud = None
    plt = None

import seaborn as sns
import sqlite3
from typing import List, Dict, Any
import asyncio
try:
    import aiohttp
except ImportError:
    aiohttp = None

# Enhanced Configuration - MUST BE FIRST, BEFORE ANYTHING ELSE
st.set_page_config(
    page_title="⚖️ AI Legal Oracle - Enterprise Edition",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': 'https://your-help-url.com',
        'Report a bug': "https://your-bug-report-url.com",
        'About': "AI Legal Oracle - The most advanced legal research assistant powered by AI"
    }
)

# --- ENSURE SESSION STATE IS ALWAYS INITIALIZED ---
def initialize_session_state():
    """Initialize all session state variables if they don't exist"""
    try:
        if 'chat_history' not in st.session_state:
            st.session_state.chat_history = []
        if 'document_processed' not in st.session_state:
            st.session_state.document_processed = False
        if 'vector_store' not in st.session_state:
            st.session_state.vector_store = None
        if 'session_id' not in st.session_state:
            st.session_state.session_id = hashlib.md5(str(datetime.now()).encode()).hexdigest()[:8]
        if 'current_mode' not in st.session_state:
            st.session_state.current_mode = "📚 Document Analysis"
        if 'error_log' not in st.session_state:
            st.session_state.error_log = []
        if 'processing_status' not in st.session_state:
            st.session_state.processing_status = "idle"
        if 'page_config_set' not in st.session_state:
            st.session_state.page_config_set = True
    except Exception as e:
        pass  # Silently fail to avoid breaking the app

# Initialize session state AFTER page config
initialize_session_state()

# Load API key
from dotenv import load_dotenv
load_dotenv()

if "OPENAI_API_KEY" not in os.environ:
    if "OPENAI_API_KEY" in st.secrets:
        os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]
    else:
        st.error("⚠️ OpenAI API key not found. Please set OPENAI_API_KEY in your environment variables or Streamlit secrets.")
        st.stop()

# Custom CSS for Professional Look
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(90deg, #1e3c72 0%, #2a5298 100%);
        padding: 2rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
    }
    .feature-card {
        background: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 5px solid #2a5298;
        margin: 1rem 0;
    }
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1rem;
        border-radius: 8px;
        color: white;
        text-align: center;
    }
    .chat-message {
        padding: 1rem;
        border-radius: 10px;
        margin: 0.5rem 0;
    }
    .user-message {
        background: #e3f2fd;
        border-left: 4px solid #2196f3;
    }
    .ai-message {
        background: #f3e5f5;
        border-left: 4px solid #9c27b0;
    }
    .stSelectbox > div > div > select {
        background-color: #f0f2f6;
    }
</style>
""", unsafe_allow_html=True)

# Database Setup for Analytics with Error Handling
def init_database():
    """Initialize database with proper error handling"""
    try:
        conn = sqlite3.connect('legal_analytics.db')
        cursor = conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS query_analytics (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp DATETIME,
                query TEXT,
                response_time FLOAT,
                document_type TEXT,
                user_session TEXT,
                confidence_score FLOAT
            )
        ''')
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        st.error(f"Database initialization error: {str(e)}")
        return False

# Initialize database safely
database_ready = init_database()

# Enhanced Session State Management with Error Handling
def initialize_session_state():
    """Initialize session state with proper error handling"""
    try:
        if 'chat_history' not in st.session_state:
            st.session_state.chat_history = []
        if 'document_processed' not in st.session_state:
            st.session_state.document_processed = False
        if 'vector_store' not in st.session_state:
            st.session_state.vector_store = None
        if 'session_id' not in st.session_state:
            st.session_state.session_id = hashlib.md5(str(datetime.now()).encode()).hexdigest()[:8]
        if 'current_mode' not in st.session_state:
            st.session_state.current_mode = "📚 Document Analysis"
        if 'error_log' not in st.session_state:
            st.session_state.error_log = []
        if 'processing_status' not in st.session_state:
            st.session_state.processing_status = "idle"
    except Exception as e:
        st.error(f"Error initializing session state: {str(e)}")


# Always ensure session state is initialized before any feature logic runs
try:
    initialize_session_state()
except Exception as e:
    st.error(f"Session state initialization failed: {str(e)}")

# Enhanced Header
st.markdown("""
<div class="main-header">
    <h1>⚖️ AI Legal Oracle - Enterprise Edition</h1>
    <p>Advanced Legal Research Assistant with Multi-Modal Intelligence, Analytics & Compliance Tools</p>
</div>
""", unsafe_allow_html=True)

# Sidebar - Advanced Features with Error Handling
with st.sidebar:
    st.header("🎛️ Control Panel")
    
    # Add status indicator
    if st.session_state.processing_status == "processing":
        st.warning("⚠️ Processing in progress...")
    elif st.session_state.processing_status == "error":
        st.error("❌ Last operation failed")
    else:
        st.success("✅ System Ready")
    
    # Mode Selection with session state tracking
    app_mode = st.selectbox(
        "🔧 Select Mode",
        ["📚 Document Analysis", "💬 Legal Chat", "📊 Analytics Dashboard", "🔍 Case Law Search", "📋 Compliance Checker"],
        key="app_mode_selector"
    )
    
    # Update session state when mode changes (no manual rerun needed)
    if 'current_mode' not in st.session_state or st.session_state.current_mode != app_mode:
        st.session_state.current_mode = app_mode
        st.session_state.processing_status = "idle"
    
    # AI Model Selection
    model_choice = st.selectbox(
        "🤖 AI Model",
        ["gpt-4o-mini (Cost Optimized)", "gpt-4 (Premium)", "gpt-3.5-turbo (Standard)"]
    )
    
    # Language Selection
    language = st.selectbox(
        "🌐 Language",
        ["English", "Spanish", "French", "German", "Chinese", "Hindi"]
    )
    
    # Legal Domain
    legal_domain = st.selectbox(
        "⚖️ Legal Domain",
        ["General", "Corporate Law", "Employment Law", "Contract Law", "IP Law", "Criminal Law", "Family Law"]
    )
    
    # Advanced Settings - Initialize default values first
    temperature = 0.1
    max_tokens = 300
    search_depth = 3
    
    with st.expander("⚙️ Advanced Settings"):
        temperature = st.slider("🌡️ Response Creativity", 0.0, 1.0, 0.1, 0.1)
        max_tokens = st.slider("📝 Max Response Length", 100, 1000, 300, 50)
        search_depth = st.slider("🔍 Search Depth", 1, 10, 3)
        
    # System Status and Quick Actions
    st.markdown("### 📈 Session Stats")
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Queries", len(st.session_state.chat_history))
    with col2:
        st.metric("Session", st.session_state.session_id)
    
    # Quick Actions
    st.markdown("### ⚡ Quick Actions")
    if st.button("🔄 Reset Session", help="Clear all data and restart", key="reset_session_btn"):
        try:
            # Clear all data except app mode selector
            keys_to_clear = [k for k in st.session_state.keys() if k != "app_mode_selector"]
            for key in keys_to_clear:
                del st.session_state[key]
            initialize_session_state()
            st.success("✅ Session reset successfully!")
            time.sleep(0.5)
            st.rerun()
        except Exception as e:
            st.error(f"Reset failed: {str(e)}")
    
    if st.button("🧹 Clear Cache", help="Clear processing cache"):
        try:
            st.cache_data.clear()
            st.cache_resource.clear()
            st.success("Cache cleared!")
        except Exception as e:
            st.error(f"Cache clear failed: {str(e)}")
    
    # Error Log (if any errors occurred)
    if st.session_state.error_log:
        with st.expander("⚠️ Error Log"):
            for error in st.session_state.error_log[-5:]:  # Show last 5 errors
                st.text(f"{error['time']}: {error['message']}")
            if st.button("Clear Error Log", key="clear_error_log_btn"):
                st.session_state.error_log = []
                st.success("✅ Error log cleared!")

# Utility Functions for Error Handling and Memory Management
def safe_operation(operation_name, operation_func, *args, **kwargs):
    """Safely execute operations with error handling and logging"""
    try:
        st.session_state.processing_status = "processing"
        result = operation_func(*args, **kwargs)
        st.session_state.processing_status = "idle"
        return result
    except Exception as e:
        st.session_state.processing_status = "error"
        error_msg = f"{operation_name} failed: {str(e)}"
        st.session_state.error_log.append({
            'time': datetime.now().strftime('%H:%M:%S'),
            'message': error_msg
        })
        st.error(error_msg)
        return None

def cleanup_memory():
    """Clean up memory and cache"""
    try:
        import gc
        gc.collect()
        st.cache_data.clear()
    except Exception as e:
        pass

@st.cache_data(ttl=3600)  # Cache for 1 hour
def cached_text_extraction(file_content, file_name, file_type):
    """Cache text extraction to avoid reprocessing same files"""
    if file_type == "pdf":
        return extract_text_from_pdf_content(file_content)
    elif file_type == "docx" and docx is not None:
        return extract_text_from_docx_content(file_content)
    else:
        return file_content.decode("utf-8"), {"filename": file_name, "type": "TXT"}

def extract_text_from_pdf_content(file_content):
    """Extract text from PDF content"""
    text = ""
    metadata = {"pages": 0, "title": "", "author": ""}
    
    try:
        from io import BytesIO
        pdf_file = BytesIO(file_content)
        reader = PdfReader(pdf_file)
        metadata["pages"] = len(reader.pages)
        
        if reader.metadata:
            metadata["title"] = reader.metadata.get("/Title", "")
            metadata["author"] = reader.metadata.get("/Author", "")
        
        for page in reader.pages:
            text += page.extract_text() + "\n"
            
    except Exception as e:
        raise Exception(f"PDF extraction failed: {str(e)}")
    
    return text, metadata

def extract_text_from_docx_content(file_content):
    """Extract text from DOCX content"""
    if docx is None:
        raise Exception("python-docx not installed")
    
    try:
        from io import BytesIO
        docx_file = BytesIO(file_content)
        doc = docx.Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")
def extract_text_from_pdf(file):
    """Enhanced PDF extraction with metadata using PyPDF2"""
    text = ""
    metadata = {"pages": 0, "title": "", "author": ""}
    
    try:
        reader = PdfReader(file)
        metadata["pages"] = len(reader.pages)
        
        # Try to get metadata
        if reader.metadata:
            metadata["title"] = reader.metadata.get("/Title", "")
            metadata["author"] = reader.metadata.get("/Author", "")
        
        for page in reader.pages:
            text += page.extract_text() + "\n"
            
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return "", metadata
    
    return text, metadata

def extract_text_from_docx(file):
    """Extract text from DOCX files"""
    if docx is None:
        st.error("python-docx not installed. Please install it to process DOCX files.")
        return ""
    
    try:
        doc = docx.Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading DOCX file: {str(e)}")
        return ""

def generate_document_summary(text: str) -> Dict[str, Any]:
    """Generate comprehensive document analysis"""
    words = text.split()
    sentences = text.split('.')
    
    # Basic statistics
    stats = {
        "word_count": len(words),
        "sentence_count": len(sentences),
        "character_count": len(text),
        "avg_sentence_length": len(words) / len(sentences) if sentences else 0
    }
    
    # Extract key legal terms
    legal_terms = []
    legal_keywords = [
        "contract", "agreement", "liability", "clause", "provision", "party",
        "defendant", "plaintiff", "jurisdiction", "compliance", "violation",
        "breach", "damages", "penalty", "regulation", "statute", "law"
    ]
    
    for term in legal_keywords:
        count = text.lower().count(term)
        if count > 0:
            legal_terms.append({"term": term, "frequency": count})
    
    stats["legal_terms"] = sorted(legal_terms, key=lambda x: x["frequency"], reverse=True)[:10]
    
    return stats

# Main Content Area with Enhanced Error Handling
required_keys = ["chat_history", "document_processed", "vector_store", "session_id", "current_mode", "error_log", "processing_status"]
for k in required_keys:
    if k not in st.session_state:
        initialize_session_state()
        break  # Only initialize once

# Ensure all variables are defined (safeguard)
try:
    _ = app_mode
    _ = model_choice
    _ = language
    _ = legal_domain
    _ = temperature
    _ = max_tokens
    _ = search_depth
except NameError:
    st.error("⚠️ Configuration error. Please refresh the page.")
    st.stop()

try:
    if app_mode == "📚 Document Analysis":
        st.header("📚 Advanced Document Analysis")
        
        # Enhanced File Upload
        col1, col2 = st.columns([2, 1])
        
        with col1:
            # Determine supported file types based on available libraries
            file_types = ['pdf', 'txt']
            if docx is not None:
                file_types.append('docx')
            
            uploaded_files = st.file_uploader(
                "📄 Upload Legal Documents",
                accept_multiple_files=True,
                type=file_types,
                help=f"Supported formats: {', '.join(file_types).upper()}. Maximum 10 files."
            )
        
        with col2:
            if uploaded_files:
                st.success(f"✅ {len(uploaded_files)} file(s) uploaded")
                
                # File preview
                for file in uploaded_files[:3]:  # Show first 3 files
                    file_size = len(file.getvalue()) / 1024  # KB
                    st.write(f"📄 {file.name} ({file_size:.1f} KB)")
        
        if uploaded_files and len(uploaded_files) <= 10:  # Limit to 10 files
            if st.button("🚀 Process Documents", type="primary", key="process_docs_btn"):
                def process_documents():
                    """Process documents with error handling"""
                    with st.spinner("🔄 Processing documents with AI..."):
                        progress_bar = st.progress(0)
                        all_text = ""
                        all_metadata = []
                        
                        for i, file in enumerate(uploaded_files):
                            try:
                                progress_bar.progress((i + 1) / len(uploaded_files))
                                
                                file_content = file.getvalue()
                                file_type = file.name.split('.')[-1].lower()
                                
                                # Use cached extraction
                                if file_type == "pdf":
                                    text, metadata = cached_text_extraction(file_content, file.name, "pdf")
                                    metadata["filename"] = file.name
                                    metadata["type"] = "PDF"
                                elif file_type == "docx" and docx is not None:
                                    text = cached_text_extraction(file_content, file.name, "docx")
                                    metadata = {"filename": file.name, "type": "DOCX"}
                                else:  # .txt
                                    text = file_content.decode("utf-8")
                                    metadata = {"filename": file.name, "type": "TXT"}
                                
                                all_text += f"\n\n--- Document: {file.name} ---\n\n" + text
                                all_metadata.append(metadata)
                                
                                # Cleanup memory after each file
                                if i % 3 == 0:  # Cleanup every 3 files
                                    cleanup_memory()
                                    
                            except Exception as e:
                                st.warning(f"⚠️ Could not process {file.name}: {str(e)}")
                                continue
                        
                        progress_bar.progress(1.0)
                        
                        if not all_text.strip():
                            raise Exception("No text could be extracted from uploaded files")
                        
                        # Enhanced text splitting with document awareness
                        splitter = RecursiveCharacterTextSplitter(
                            chunk_size=1200,  # Reduced for stability
                            chunk_overlap=150,
                            separators=["\n\n--- Document:", "\n\n", "\n", ".", "!", "?", ",", " ", ""]
                        )
                        chunks = splitter.split_text(all_text)
                        
                        # Limit chunks for stability
                        if len(chunks) > 80:  # Reduced limit
                            chunks = chunks[:80]
                            st.warning(f"⚠️ Limited to 80 chunks for optimal performance.")
                        
                        # Create embeddings with progress tracking
                        embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
                        
                        with st.spinner("🧠 Creating AI knowledge base..."):
                            vector_store = FAISS.from_texts(chunks, embedding=embeddings)
                            st.session_state.vector_store = vector_store
                            st.session_state.document_processed = True
                        
                        return chunks, all_metadata
                
                # Execute document processing safely
                result = safe_operation("Document Processing", process_documents)
                
                if result:
                    chunks, all_metadata = result
                    
                    # Document Analysis Dashboard
                    st.success("✅ Documents processed successfully!")
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    total_words = sum([len(chunk.split()) for chunk in chunks[:5]])
                    
                    with col1:
                        st.metric("📄 Documents", len(uploaded_files))
                    with col2:
                        st.metric("🧩 Text Chunks", len(chunks))
                    with col3:
                        st.metric("📝 Total Words", f"{total_words:,}")
                    with col4:
                        st.metric("💾 Vector Store", "Ready")
                    
        elif len(uploaded_files) > 10:
            st.error("⚠️ Too many files! Please upload maximum 10 files at once.")

    elif app_mode == "💬 Legal Chat":
        try:
            st.header("💬 AI Legal Assistant Chat")
            
            # Double-check vector store exists
            if not st.session_state.document_processed or st.session_state.vector_store is None:
                st.warning("⚠️ Please process documents first in the Document Analysis mode.")
                st.info("👆 Switch to 'Document Analysis' mode using the control panel to upload and process your documents.")
            else:
                # Enhanced Chat Interface with error handling
                st.markdown("### 🤖 Chat with your legal documents")
                
                # Quick suggestion buttons
                st.markdown("**💡 Quick Questions:**")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    if st.button("📋 Summarize key points", key="chat_summarize_btn"):
                        st.session_state.temp_query = "Please provide a comprehensive summary of the key legal points in these documents."
                
                with col2:
                    if st.button("⚠️ Identify risks", key="chat_risks_btn"):
                        st.session_state.temp_query = "What are the main legal risks and liabilities mentioned in these documents?"
                
                with col3:
                    if st.button("✅ Check compliance", key="chat_compliance_btn"):
                        st.session_state.temp_query = "Are there any compliance requirements or regulatory obligations mentioned?"
                
                # Chat input with validation
                user_query = st.text_input(
                    "💭 Ask anything about your legal documents:",
                    value=st.session_state.get('temp_query', ''),
                    placeholder="e.g., What are the termination clauses in this contract?",
                    max_chars=500  # Limit query length
                )
                
                if 'temp_query' in st.session_state:
                    del st.session_state.temp_query
                
                if user_query and len(user_query.strip()) > 5:
                    def process_chat_query():
                        """Process chat query with error handling"""
                        # Input validation
                        if len(user_query) > 500:
                            raise Exception("Query too long. Please keep under 500 characters.")
                        
                        # Advanced prompt template for legal context
                        legal_prompt = PromptTemplate(
                            template="""You are an expert legal AI assistant. Based on the provided legal documents, answer the user's question with:

1. **Direct Answer**: Provide a clear, concise response
2. **Legal Context**: Explain relevant legal principles
3. **Document References**: Cite specific sections when possible
4. **Risk Assessment**: Highlight any potential legal risks
5. **Recommendations**: Suggest next steps if applicable

Legal Domain: {legal_domain}
Language: {language}

Question: {question}

Context from documents:
{context}

Please provide a comprehensive legal analysis:""",
                            input_variables=["legal_domain", "language", "question", "context"]
                        )
                        
                        # Create enhanced QA chain
                        retriever = st.session_state.vector_store.as_retriever(
                            search_type="mmr",
                            search_kwargs={"k": min(search_depth, 3), "fetch_k": min(search_depth * 2, 6)}
                        )
                        
                        # Configure LLM based on user selection
                        model_map = {
                            "gpt-4o-mini (Cost Optimized)": "gpt-4o-mini",
                            "gpt-4 (Premium)": "gpt-4",
                            "gpt-3.5-turbo (Standard)": "gpt-3.5-turbo"
                        }
                        
                        llm = ChatOpenAI(
                            model_name=model_map[model_choice],
                            temperature=temperature,
                            max_tokens=min(max_tokens, 500),  # Cap tokens for stability
                            timeout=30  # 30 second timeout
                        )
                        
                        qa_chain = RetrievalQA.from_chain_type(
                            llm=llm,
                            chain_type="stuff",
                            retriever=retriever,
                            return_source_documents=True,
                            chain_type_kwargs={"prompt": legal_prompt}
                        )
                        
                        start_time = time.time()
                        
                        result = qa_chain({
                            "query": user_query,
                            "legal_domain": legal_domain,
                            "language": language
                        })
                        
                        response_time = time.time() - start_time
                        
                        # Add to chat history
                        st.session_state.chat_history.append({
                            "timestamp": datetime.now(),
                            "query": user_query,
                            "response": result["result"],
                            "sources": result["source_documents"],
                            "response_time": response_time
                        })
                        
                        # Log analytics safely
                        try:
                            if database_ready:
                                conn = sqlite3.connect('legal_analytics.db')
                                cursor = conn.cursor()
                                cursor.execute('''
                                    INSERT INTO query_analytics 
                                    (timestamp, query, response_time, document_type, user_session, confidence_score)
                                    VALUES (?, ?, ?, ?, ?, ?)
                                ''', (
                                    datetime.now(),
                                    user_query,
                                    response_time,
                                    legal_domain,
                                    st.session_state.session_id,
                                    0.85  # Default confidence score
                                ))
                                conn.commit()
                                conn.close()
                        except Exception as db_error:
                            # Don't break the app if database logging fails
                            pass
                        
                        return result, response_time
                    
                    # Execute chat query safely
                    with st.spinner("🧠 AI is analyzing your legal documents..."):
                        chat_result = safe_operation("Chat Query", process_chat_query)
                    
                    if chat_result:
                        result, response_time = chat_result
                        
                        # Display response
                        st.success("✅ Analysis complete!")
                        
                        # Show response
                        st.markdown("### 🧾 Answer:")
                        st.write(result["result"])
                        
                        # Show cost estimate
                        input_tokens = len(user_query.split()) * 1.3
                        output_tokens = len(result["result"].split()) * 1.3
                        estimated_cost = (input_tokens * 0.00015 + output_tokens * 0.0006) / 1000
                        
                        st.caption(f"💸 Response time: {response_time:.2f}s | Estimated cost: ~${estimated_cost:.6f}")
                        
                        # Show sources
                        with st.expander("📚 Sources & References"):
                            for j, doc in enumerate(result["source_documents"], 1):
                                st.markdown(f"**📄 Source {j}:**")
                                st.text(doc.page_content[:300] + "...")
                                if j < len(result["source_documents"]):
                                    st.markdown("---")
                
                elif user_query and len(user_query.strip()) <= 5:
                    st.warning("⚠️ Please enter a more detailed question (at least 6 characters).")
                
                # Display chat history (last 3 conversations)
                if st.session_state.chat_history:
                    st.markdown("### 💬 Recent Conversations")
                    
                    for i, chat in enumerate(reversed(st.session_state.chat_history[-3:])):
                        with st.container():
                            st.markdown(f"""
                            <div class="chat-message user-message">
                                <strong>👤 You ({chat['timestamp'].strftime('%H:%M')}):</strong><br>
                                {chat['query']}
                            </div>
                            """, unsafe_allow_html=True)
                            
                            st.markdown(f"""
                            <div class="chat-message ai-message">
                                <strong>🤖 AI Legal Assistant:</strong><br>
                                {chat['response'][:500]}{'...' if len(chat['response']) > 500 else ''}
                            </div>
                            """, unsafe_allow_html=True)
                            
                            st.markdown("---")
                            
        except Exception as e:
            st.error(f"Chat mode error: {str(e)}")
            st.info("💡 Try refreshing the page or switching to a different mode.")

    elif app_mode == "📊 Analytics Dashboard":
        try:
            st.header("📊 Legal Analytics Dashboard")
            
            if not database_ready:
                st.error("⚠️ Analytics database is not available. Some features may not work.")
            
            # Load analytics data safely
            try:
                conn = sqlite3.connect('legal_analytics.db')
                df = pd.read_sql_query("SELECT * FROM query_analytics ORDER BY timestamp DESC LIMIT 1000", conn)
                conn.close()
                
                if not df.empty:
                    # Overview metrics
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        st.metric("📊 Total Queries", len(df))
                    with col2:
                        avg_response_time = df['response_time'].mean()
                        st.metric("⚡ Avg Response Time", f"{avg_response_time:.2f}s")
                    with col3:
                        unique_sessions = df['user_session'].nunique()
                        st.metric("👥 Unique Sessions", unique_sessions)
                    with col4:
                        avg_confidence = df['confidence_score'].mean()
                        st.metric("🎯 Avg Confidence", f"{avg_confidence:.2%}")
                    
                    # Charts with error handling
                    tab1, tab2, tab3 = st.tabs(["📈 Usage Trends", "🏷️ Query Categories", "⚡ Performance"])
                    
                    with tab1:
                        try:
                            # Daily usage trend
                            df['date'] = pd.to_datetime(df['timestamp']).dt.date
                            daily_queries = df.groupby('date').size().reset_index(name='queries')
                            
                            if len(daily_queries) > 1:
                                fig = px.line(
                                    daily_queries, 
                                    x='date', 
                                    y='queries',
                                    title="Daily Query Volume",
                                    markers=True
                                )
                                st.plotly_chart(fig, use_container_width=True)
                            else:
                                st.info("📊 Not enough data for trend analysis yet.")
                        except Exception as e:
                            st.error(f"Error creating usage trends: {str(e)}")
                    
                    with tab2:
                        try:
                            # Query categories
                            domain_counts = df['document_type'].value_counts()
                            
                            if len(domain_counts) > 0:
                                fig = px.bar(
                                    x=domain_counts.index,
                                    y=domain_counts.values,
                                    title="Queries by Legal Domain",
                                    labels={'x': 'Legal Domain', 'y': 'Number of Queries'}
                                )
                                st.plotly_chart(fig, use_container_width=True)
                            else:
                                st.info("📊 No domain data available yet.")
                        except Exception as e:
                            st.error(f"Error creating category chart: {str(e)}")
                    
                    with tab3:
                        try:
                            # Performance metrics
                            if len(df) > 5:
                                fig = px.scatter(
                                    df.head(100),  # Limit to last 100 for performance
                                    x='response_time',
                                    y='confidence_score',
                                    title="Response Time vs Confidence Score",
                                    labels={'response_time': 'Response Time (seconds)', 'confidence_score': 'Confidence Score'}
                                )
                                st.plotly_chart(fig, use_container_width=True)
                            else:
                                st.info("📊 Need at least 5 queries for performance analysis.")
                        except Exception as e:
                            st.error(f"Error creating performance chart: {str(e)}")
                    
                    # Recent queries table
                    st.subheader("📋 Recent Queries")
                    recent_queries = df.head(10)[['timestamp', 'query', 'response_time', 'document_type']]
                    recent_queries['query'] = recent_queries['query'].str[:100] + '...'  # Truncate for display
                    st.dataframe(recent_queries, use_container_width=True)
                    
                else:
                    st.info("📊 No analytics data available yet. Start asking questions to see analytics!")
            
            except Exception as e:
                st.error(f"Error loading analytics: {str(e)}")
                st.info("💡 Try using the app more to generate analytics data.")
                
        except Exception as e:
            st.error(f"Analytics dashboard error: {str(e)}")

    elif app_mode == "🔍 Case Law Search":
        st.header("🔍 Legal Case Law & Precedent Search")
        st.info("🚧 **Coming Soon**: Advanced case law search with precedent analysis, citation tracking, and legal reasoning chains.")
        
        # Mock interface for demonstration
        search_query = st.text_input("🔍 Search legal cases and precedents:")
        
        if search_query:
            with st.spinner("🔍 Searching legal databases..."):
                time.sleep(1)  # Simulate search
                
                st.success("🎯 Feature in development - Mock results shown")
                
                # Mock results
                mock_cases = [
                    {"title": "Smith v. Johnson Corp", "year": "2023", "court": "9th Circuit", "relevance": "95%"},
                    {"title": "Privacy Rights Coalition v. Tech Co", "year": "2022", "court": "Supreme Court", "relevance": "89%"},
                    {"title": "Employment Standards Board v. ABC Inc", "year": "2021", "court": "District Court", "relevance": "84%"}
                ]
                
                for case in mock_cases:
                    with st.expander(f"📚 {case['title']} ({case['year']})"):
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.write(f"**Court:** {case['court']}")
                        with col2:
                            st.write(f"**Year:** {case['year']}")
                        with col3:
                            st.write(f"**Relevance:** {case['relevance']}")
                        
                        st.write("**Summary:** Lorem ipsum legal case summary with key holdings and precedential value...")

    elif app_mode == "📋 Compliance Checker":
        st.header("📋 Legal Compliance & Risk Assessment")
        st.info("🚧 **Coming Soon**: Automated compliance checking against regulatory frameworks, risk scoring, and compliance gap analysis.")
        
        # Mock interface
        compliance_framework = st.selectbox(
            "📋 Select Compliance Framework",
            ["GDPR", "SOX", "HIPAA", "CCPA", "ISO 27001", "Custom Framework"]
        )
        
        if st.button("🔍 Run Compliance Check"):
            with st.spinner("🔍 Analyzing documents for compliance..."):
                time.sleep(2)  # Simulate analysis
                
                st.success("✅ Compliance analysis complete! (Demo Mode)")
                
                # Mock compliance results
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.metric("📊 Compliance Score", "78%", "📈 +5%")
                with col2:
                    st.metric("⚠️ Issues Found", "12", "📉 -3")
                with col3:
                    st.metric("🔴 Critical Issues", "2", "➡️ 0")

except Exception as main_error:
    st.error(f"🚨 Application Error: {str(main_error)}")
    st.info("💡 Try switching to a different mode or refreshing the page.")
    
    # Add error to log safely
    try:
        if 'error_log' in st.session_state:
            st.session_state.error_log.append({
                'time': datetime.now().strftime('%H:%M:%S'),
                'message': f"Main app error: {str(main_error)}"
            })
    except:
        pass
    
    # Provide recovery options without aggressive rerun
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔄 Refresh Page", key="error_refresh_btn"):
            time.sleep(0.5)
            st.rerun()
    with col2:
        if st.button("🆘 Reset Application", key="error_reset_btn"):
            try:
                for key in list(st.session_state.keys()):
                    if key not in ["app_mode_selector"]:
                        del st.session_state[key]
                initialize_session_state()
                time.sleep(0.5)
                st.rerun()
            except:
                st.rerun()

# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #666; padding: 2rem 0;">
    <p>⚖️ <strong>AI Legal Oracle Enterprise Edition</strong> | Powered by Advanced AI & Machine Learning</p>
    <p>🔒 Your documents are processed securely and never stored permanently | 💡 Built with Streamlit & LangChain</p>
</div>
""", unsafe_allow_html=True)